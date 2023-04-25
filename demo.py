from tkinter import *
from tkinter import messagebox
from PIL import Image, ImageTk
from tkinter import filedialog
import mysql.connector
from tkinter import ttk
import os
import time
import docx
from docx import Document
from docx.shared import Inches
from docx2pdf import convert


class student_class:
    def _init_(self, mm):
        self.mm = mm
        self.mm.geometry('810x603+200+85')
        self.mm.title("User Login")
        self.mm.resizable('false', 'false')
        self.mm.configure(bg="black")

        self.reg_id = StringVar()
        self.name = StringVar()
        self.fname = StringVar()
        self.gender = StringVar()
        self.dob = StringVar()
        self.email = StringVar()
        self.phone = StringVar()
        self.address = StringVar()
        self.class_id = StringVar()
        self.img = StringVar()
        self.search = StringVar()
        self.select = StringVar()
        self.class_select = StringVar()

        # =============== Title Frame  ===================
        title_frame = Frame(self.mm, width=810, height=50, relief=RIDGE, bd=4, bg="#00ffff")
        title_frame.place(x=0, y=0)

        # =============== Title Student Registration ===================
        lbl_title = Label(title_frame, text="Student Registration Form", font=("Times", 25, "bold"),
                          bg="#00ffff").place(x=220, y=0)

        # ================ Search Frame ===============================
        search_frame = LabelFrame(self.mm, text="Search Student Information", font=("arial", 11, "bold"),
                                  width=810, height=65, relief=RIDGE, bd=4, fg="red", bg="#FFFFE0")
        search_frame.place(x=0, y=280)

        # =========== Search Frame Labels ==============
        lbl_select_class = Label(search_frame, text="Select Class",
                                 font=("arial", 12, "bold"), bg="seashell").place(x=0, y=3)

        combo_select_class = ttk.Combobox(search_frame, state="readonly", textvariable=self.class_select,
                                          font=("arial", 10, "bold"), width=12)
        combo_select_class["value"] = ("Select Class", "Nursery", "Prep 1", "Prep 2", "Class 1",
                                       "Class 2", "Class 3", "Class 4", "Class 5", "Class 6", "Class 8", "Class 9",
                                       "Matriculation")
        combo_select_class.current(0)
        combo_select_class.place(x=110, y=5)

        lbl_select_class = Label(search_frame, text="Search By",
                                 font=("arial", 12, "bold"), bg="seashell").place(x=220, y=3)

        combo_select_class = ttk.Combobox(search_frame, textvariable=self.select,
                                          state="readonly", font=("arial", 10, "bold"), width=14)
        combo_select_class["value"] = ("Select By", "ID", "Std_Name")
        combo_select_class.current(0)
        combo_select_class.place(x=310, y=5)

        entry_search = Entry(search_frame, font=("arial", 12, "bold"), width=14, textvariable=self.search).place(x=440,
                                                                                                                 y=5)
        btn_search = Button(search_frame, text="Search", font=("arial", 11, "bold"), bg="#00ffff", width=10,
                            command=self.search_data).place(x=580, y=0)
        btn_view_all = Button(search_frame, text="View All", font=("arial", 11, "bold"), bg="#00ffff", width=10,
                              command=self.show_data).place(x=690, y=0)
        try:
            treeview_frame = LabelFrame(self.mm, text="Student Details", font=("arial", 11, "bold"), relief=RIDGE, bd=4,
                                        bg="#FFE4B5", fg="red")
            treeview_frame.place(x=0, y=343, width=810, height=200)

            scroll_x = ttk.Scrollbar(treeview_frame, orient=HORIZONTAL)
            scroll_y = ttk.Scrollbar(treeview_frame, orient=VERTICAL)
            style = ttk.Style()
            style.theme_use("clam")
            style.configure("Treeview", background="#FFE4E1", forebackground="yellow",
                            fieldbackground="#FFE4E1")
            style.map("Treeview", background=[('selected', '#074463')])
            self.std_table = ttk.Treeview(treeview_frame, columns=("ID", "C", "N", "FN", "G", "D", "P", "E", "A", "I"),
                                          xscrollcommand=scroll_x.set, yscrollcommand=scroll_y.set)

            scroll_x.pack(side=BOTTOM, fill=X)
            scroll_y.pack(side=RIGHT, fill=Y)

            scroll_x.config(command=self.std_table.xview)
            scroll_y.config(command=self.std_table.yview)

            self.std_table.heading("ID", text="Reg. ID")
            self.std_table.heading("C", text="Class")
            self.std_table.heading("N", text="Name")
            self.std_table.heading("FN", text="Father Name")
            self.std_table.heading("G", text="Gender")
            self.std_table.heading("D", text="D.O.B")
            self.std_table.heading("P", text="Email")
            self.std_table.heading("E", text="Phone")
            self.std_table.heading("A", text="Address")
            self.std_table.heading("I", text="Image")
            self.std_table["show"] = "headings"

            self.std_table.column("ID", width=50)
            self.std_table.column("N", width=120)
            self.std_table.column("FN", width=120)
            self.std_table.column("G", width=60)
            self.std_table.column("D", width=80)
            self.std_table.column("P", width=100)
            self.std_table.column("C", width=80)
            self.std_table.pack(fill=BOTH, expand=1)
            self.std_table.bind("<ButtonRelease>", self.get_cursor)
            self.show_data()
        except Exception as e:
            messagebox.showerror("Error", f"due to {e}")

        self.img_frame = LabelFrame(self.mm, text="Student Information", font=("arial", 11, "bold"),
                                    width=810, height=230, fg="red", relief=RIDGE, bd=4, bg="seashell")
        self.img_frame.place(x=0, y=50)

        self.photo_frame = Frame(self.img_frame, width=200, height=190, relief=RIDGE, bd=3, bg="white")
        self.photo_frame.place(x=5, y=7)

        lbl_std_id = Label(self.img_frame, text="Std. Reg. No", font=("arial", 12, "bold"), bg="seashell").place(x=220,
                                                                                                                 y=20)
        lbl_std_name = Label(self.img_frame, text="Std.  Name", font=("arial", 12, "bold"), bg="seashell").place(x=220,
                                                                                                                 y=47)
        lbl_father_name = Label(self.img_frame, text="Father Name", font=("arial", 12, "bold"), bg="seashell").place(
            x=220, y=77)
        lbl_std_gender = Label(self.img_frame, text="Std. Gender", font=("arial", 12, "bold"), bg="seashell").place(
            x=220, y=107)
        lbl_std_dob = Label(self.img_frame, text="Std. D.O.B", font=("arial", 12, "bold"), bg="seashell").place(x=220,
                                                                                                                y=137)
        lbl_std_email = Label(self.img_frame, text="Std. Email", font=("arial", 12, "bold"), bg="seashell").place(x=220,
                                                                                                                  y=167)

        lbl_std_phone = Label(self.img_frame, text="Std. Phone", font=("arial", 12, "bold"), bg="seashell").place(x=520,
                                                                                                                  y=20)
        lbl_std_class = Label(self.img_frame, text="Std. Class", font=("arial", 12, "bold"), bg="seashell").place(x=520,
                                                                                                                  y=47)
        lbl_std_address = Label(self.img_frame, text="Address", font=("arial", 12, "bold"), bg="seashell").place(x=520,
                                                                                                                 y=77)
        lbl_std_upload = Label(self.img_frame, text="Upload", font=("arial", 12, "bold"), bg="seashell").place(x=520,
                                                                                                               y=107)

        std_id = Entry(self.img_frame, textvariable=self.reg_id, font=("arial", 12, "bold"), width=15,
                       bg="seashell").place(x=350, y=20)

        std_name = Entry(self.img_frame, textvariable=self.name, font=("arial", 12, "bold"), width=15,
                         bg="seashell").place(x=350, y=47)

        std_fname = Entry(self.img_frame, width=15, textvariable=self.fname, font=("arial", 12, "bold")).place(x=350,
                                                                                                               y=77)

        combo_gender = ttk.Combobox(self.img_frame, textvariable=self.gender, state="readonly",
                                    font=("arial", 11, "bold"), width=15)
        combo_gender["value"] = ("Select Gender", "Male", "Female")
        combo_gender.current(0)
        combo_gender.place(x=350, y=107)

        std_dob = Entry(self.img_frame, width=15, textvariable=self.dob, font=("arial", 12, "bold"),
                        bg="seashell").place(x=350, y=137)

        std_email = Entry(self.img_frame, width=15, textvariable=self.email, font=("arial", 12, "bold"),
                          bg="seashell").place(x=350, y=167)

        std_phone = Entry(self.img_frame, width=15, textvariable=self.phone, font=("arial", 12, "bold"),
                          bg="seashell").place(x=625, y=20)

        combo_class = ttk.Combobox(self.img_frame, textvariable=self.class_id,
                                   state="readonly", font=("arial", 11, "bold"), width=15)
        combo_class["value"] = ("Select Class", "Nursery", "Class 1",
                                "Class 2", "Class 3", "Class 4", "Class 5", "Class 6", "Class 8", "Class 9",
                                "Matriculation")
        combo_class.current(0)
        combo_class.place(x=625, y=47)

        std_address = Entry(self.img_frame, width=15, textvariable=self.address,
                            font=("arial", 12, "bold"), bg="seashell").place(x=625, y=77)

        std_img = Button(self.img_frame, text="Upload Image", command=self.upload_img, width=16,
                         font=("arial", 10, "bold")).place(x=625, y=107)

        self.photo = Entry(self.img_frame, width=40, textvariable=self.img, bg="seashell")
        self.photo.place(x=520, y=142)

        self.btn_frame = LabelFrame(self.mm, text="Buttons", font=("arial", 11, "bold"),
                                    width=810, height=62, fg="red", relief=RIDGE, bd=4, bg="white")
        self.btn_frame.place(x=0, y=540)

        btn_submit = Button(self.btn_frame, width=12, command=self.submit, text="Submit", font=("arial", 11, "bold"),
                            bg="#00ffff").place(x=30, y=2)
        btn_update = Button(self.btn_frame, width=12, command=self.update_data, text="Update",
                            font=("arial", 11, "bold"), bg="#00ffff").place(x=150, y=2)
        btn_delete = Button(self.btn_frame, width=12, command=self.delete_data, text="Delete",
                            font=("arial", 11, "bold"), bg="#00ffff").place(x=270, y=2)
        btn_reset = Button(self.btn_frame, width=12, command=self.clear, text="Clear", font=("arial", 11, "bold"),
                           bg="#00ffff").place(x=390, y=2)

        self.download_icon = ImageTk.PhotoImage(file="download_icon.png")
        self.lbl_logo_download = Button(self.btn_frame, image=self.download_icon, command=self.word)
        self.lbl_logo_download.place(x=700, y=0)

    def upload_img(self):
        FilePath = filedialog.askopenfilename(initialdir=os.getcwd(), title="open images",
                                              filetype=(("JPG File", ".jpg"), ("PNG File", ".png"), ("All Files", ".")))
        self.photo.insert(0, FilePath)
        self.ig = Image.open(FilePath)
        new_pic = self.ig.resize((185, 175), Image.ANTIALIAS)
        self.ig = ImageTk.PhotoImage(new_pic)
        self.img_lbl = Label(self.photo_frame, image=self.ig, compound=TOP, width=190, height=180, bg="white")
        self.img_lbl.place(x=0, y=0)

    def submit(self):
        if self.reg_id.get() == "" or self.name.get() == "" or self.gender.get() == "":
            messagebox.showerror("Error", "All fields are required")
        else:
            try:
                conn = mysql.connector.connect(host="localhost", user="root", database="sms")
                cursor = conn.cursor()
                cursor.execute("insert into student_data values(%s ,%s ,%s ,%s ,%s ,%s, %s, %s ,%s ,%s)", (
                    self.reg_id.get(),
                    self.class_id.get(),
                    self.name.get(),
                    self.fname.get(),
                    self.gender.get(),
                    self.dob.get(),
                    self.email.get(),
                    self.phone.get(),
                    self.address.get(),
                    self.img.get(),

                ))
                conn.commit()
                self.show_data()
                conn.close()
                messagebox.showinfo("Success", f"Student Has been added", parent=self.mm)
            except Exception as e:
                messagebox.showerror("Error", f"Due To:{e}", parent=self.mm)

    def show_data(self):
        conn = mysql.connector.connect(host="localhost", user="root", database="sms")
        cursor = conn.cursor()
        cursor.execute("select * from student_data")
        data = cursor.fetchall()
        if len(data) != 0:
            self.std_table.delete(*self.std_table.get_children())
            for i in data:
                self.std_table.insert("", END, values=i)
                conn.commit()
            conn.close()

    def get_cursor(self, event=""):
        row = self.std_table.focus()
        content = self.std_table.item(row)
        data = content["values"]
        self.reg_id.set(data[0])
        self.class_id.set(data[1])
        self.name.set(data[2])
        self.fname.set(data[3])
        self.gender.set(data[4])
        self.dob.set(data[5])
        self.phone.set(data[7])
        self.email.set(data[6])
        self.address.set(data[8])
        self.img.set(data[9])
        a = self.img.get()
        self.ig = Image.open(a)
        new_pic = self.ig.resize((185, 175), Image.ANTIALIAS)
        self.ig = ImageTk.PhotoImage(new_pic)
        self.img_lbl = Label(self.photo_frame, image=self.ig, compound=TOP, width=190, height=180, bg="white")
        self.img_lbl.place(x=0, y=0)

    def update_data(self):
        global conn
        if self.reg_id.get() == "" or self.name.get() == "" or self.photo.get() == "":
            messagebox.showerror("Error", "All fields are required")
        else:
            try:
                update = messagebox.askyesno("Update", "Are You Sure Update this Student data", parent=self.mm)
                if update > 0:
                    conn = mysql.connector.connect(host="localhost", user="root", database="sms")
                    cursor = conn.cursor()
                    cursor.execute(
                        "update student_data set Class=%s,Std_Name=%s,Father_Name=%s,Gender=%s,DOB=%s,Email=%s,Phone=%s,Address=%s,Image=%s where ID=%s",
                        (
                            self.class_id.get(),
                            self.name.get(),
                            self.fname.get(),
                            self.gender.get(),
                            self.dob.get(),
                            self.email.get(),
                            self.phone.get(),
                            self.address.get(),
                            self.img.get(),
                            self.reg_id.get(),
                        ))
                else:
                    if not update:
                        return
                conn.commit()
                self.show_data()
                conn.close()

                messagebox.showinfo("Success", "Student Successfully Updated", parent=self.mm)
            except Exception as e:
                messagebox.showerror("Error", f"Due To:{e}", parent=self.mm)

    def delete_data(self):
        global conn
        if self.reg_id.get() == "" or self.name.get() == "" or self.photo.get() == "":
            messagebox.showerror("Error", "All fields are required")
        else:
            try:
                delete = messagebox.askyesno("Update", "Are You Sure Update this Student data", parent=self.mm)
                if delete > 0:
                    conn = mysql.connector.connect(host="localhost", user="root", database="sms")
                    cursor = conn.cursor()
                    query = ("delete from student_data where ID=%s")
                    value = (self.reg_id.get(),)
                    cursor.execute(query, value)
                else:
                    if not delete:
                        return
                    conn.commit()
                    self.show_data()
                    conn.close()

                    messagebox.showinfo("Success", "Student Successfully Deleted", parent=self.mm)
            except Exception as e:
                messagebox.showerror("Error", f"Due To:{e}", parent=self.mm)

    def clear(self):
        self.reg_id.set("")
        self.class_id.set("")
        self.name.set("")
        self.fname.set("")
        self.gender.set("")
        self.dob.set("")
        self.email.set("")
        self.phone.set("")
        self.address.set("")
        self.img.set("")

    def search_data(self):
        global conn
        if self.search.get == "":
            messagebox.showerror("Error", "Please Select Option")
        else:
            try:
                conn = mysql.connector.connect(host="localhost", user="root", database="sms")
                cursor = conn.cursor()
                cursor.execute("select * from student_data where Class LIKE '%" + str(
                    self.class_select.get()) + "%'" + " and " + str(self.select.get()) + " LIKE '%" + str(
                    self.search.get()) + "%'")
                data = cursor.fetchall()
                if len(data) != 0:
                    self.std_table.delete(*self.std_table.get_children())
                    for i in data:
                        self.std_table.insert("", END, values=i)
                        conn.commit()

                else:
                    messagebox.showerror("Error", "Student Not Found")
                conn.close()
            except Exception as e:
                messagebox.showerror("Error", f"Due To:{e}", parent=self.mm)

    def slider(self):
        self.x -= 1
        if self.x == 0:
            self.x = 500
            time.sleep(2)
            self.new_img = self.m1
            self.m1 = self.m2
            self.m2 = self.m3
            self.m3 = self.new_img
            self.m1_lbl.config(image=self.m2)
            self.m2_lbl.config(image=self.m1)
            self.m3_lbl.config(image=self.new_img)

        self.m3_lbl.place(x=self.x, y=0)
        self.m3_lbl.after(4, self.slider)

    def word(self):
        doc = docx.Document()
        doc.add_picture(self.img.get(), width=docx.shared.Inches(2), height=docx.shared.Inches(2))
        doc.add_paragraph("Student Registration No  :-  " + self.reg_id.get())
        doc.add_paragraph("Student Class                      :-  " + self.class_id.get())
        doc.add_paragraph("Student Name                   :-  " + self.name.get())
        doc.add_paragraph("Student Father Name     :-  " + self.fname.get())
        doc.add_paragraph("Student Gender          :-  " + self.gender.get())
        doc.add_paragraph("Student D.O.B           :-  " + self.dob.get())
        doc.add_paragraph("Student Email           :-  " + self.email.get())
        doc.add_paragraph("Student Phone No        :-  " + self.phone.get())
        doc.add_paragraph("Student Address         :-  " + self.address.get())
        doc.save("student{}.docx".format(self.reg_id.get()))
        convert("student{}.docx".format(self.reg_id.get()),
                r"D:\Mateen Mureed\VS CODE\Project\student{}.pdf".format(self.reg_id.get()))


if _name_ == '_main_':
    mm = Tk()
    obj = student_class(mm)
    mm.mainloop()